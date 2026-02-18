"""Export service for transcript formats.

Supports: TXT, SRT, VTT, DOCX, PDF
"""

import io
import logging
from dataclasses import dataclass
from datetime import timedelta
from typing import Any

logger = logging.getLogger(__name__)


@dataclass
class ExportSegment:
    """Segment data for export."""

    index: int
    start_time: float
    end_time: float
    text: str
    speaker: str | None = None
    speaker_name: str | None = None
    edited: bool = False
    highlight_color: str | None = None
    comments: list[str] | None = None


@dataclass
class SpeakerStats:
    """Per-speaker statistics."""

    speaker_name: str
    word_count: int
    speaking_time: float  # seconds
    word_percent: float   # 0-100
    time_percent: float   # 0-100


@dataclass
class ExportData:
    """Data structure for transcript export."""

    title: str
    language: str | None
    model_used: str | None
    word_count: int | None
    duration_seconds: float | None
    segments: list[ExportSegment]
    speakers: dict[str, str]  # speaker_label -> speaker_name
    speaker_stats: list[SpeakerStats] | None = None
    ai_summary: dict | None = None


def compute_speaker_stats(
    segments: list[ExportSegment],
    speakers: dict[str, str],
) -> list[SpeakerStats]:
    """Compute per-speaker word count and speaking time from segments."""
    from collections import defaultdict

    words_by_speaker: dict[str, int] = defaultdict(int)
    time_by_speaker: dict[str, float] = defaultdict(float)

    for seg in segments:
        name = speakers.get(seg.speaker or "", seg.speaker or "Unknown")
        words_by_speaker[name] += len(seg.text.split())
        time_by_speaker[name] += max(0, seg.end_time - seg.start_time)

    total_words = sum(words_by_speaker.values()) or 1
    total_time = sum(time_by_speaker.values()) or 1.0

    stats = []
    for name in speakers.values():
        if name in words_by_speaker:
            stats.append(SpeakerStats(
                speaker_name=name,
                word_count=words_by_speaker[name],
                speaking_time=time_by_speaker[name],
                word_percent=round(words_by_speaker[name] / total_words * 100, 1),
                time_percent=round(time_by_speaker[name] / total_time * 100, 1),
            ))

    return sorted(stats, key=lambda s: s.word_count, reverse=True)


def format_timestamp_srt(seconds: float) -> str:
    """Format seconds to SRT timestamp (HH:MM:SS,mmm)."""
    td = timedelta(seconds=seconds)
    hours, remainder = divmod(int(td.total_seconds()), 3600)
    minutes, secs = divmod(remainder, 60)
    milliseconds = int((seconds % 1) * 1000)
    return f"{hours:02d}:{minutes:02d}:{secs:02d},{milliseconds:03d}"


def format_timestamp_vtt(seconds: float) -> str:
    """Format seconds to VTT timestamp (HH:MM:SS.mmm)."""
    td = timedelta(seconds=seconds)
    hours, remainder = divmod(int(td.total_seconds()), 3600)
    minutes, secs = divmod(remainder, 60)
    milliseconds = int((seconds % 1) * 1000)
    return f"{hours:02d}:{minutes:02d}:{secs:02d}.{milliseconds:03d}"


def format_timestamp_readable(seconds: float) -> str:
    """Format seconds to readable timestamp (MM:SS or HH:MM:SS)."""
    td = timedelta(seconds=seconds)
    total_seconds = int(td.total_seconds())
    hours, remainder = divmod(total_seconds, 3600)
    minutes, secs = divmod(remainder, 60)
    if hours > 0:
        return f"{hours:d}:{minutes:02d}:{secs:02d}"
    return f"{minutes:d}:{secs:02d}"


class ExportService:
    """Service for exporting transcripts to various formats."""

    def export_txt(self, data: ExportData, include_timestamps: bool = True) -> str:
        """Export transcript to plain text format.

        Args:
            data: Export data
            include_timestamps: Whether to include timestamps

        Returns:
            Plain text content
        """
        lines = []

        # Header
        lines.append(f"Transcript: {data.title}")
        lines.append("=" * 50)
        if data.language:
            lines.append(f"Language: {data.language.upper()}")
        if data.word_count:
            lines.append(f"Word Count: {data.word_count:,}")
        if data.duration_seconds:
            lines.append(f"Duration: {format_timestamp_readable(data.duration_seconds)}")
        lines.append("")
        lines.append("-" * 50)
        lines.append("")

        # Speaker statistics
        if data.speaker_stats:
            lines.append("Speaker Statistics")
            lines.append("-" * 50)
            lines.append(f"{'Speaker':<20} {'Words':>8} {'Time':>10} {'%':>6}")
            lines.append("-" * 50)
            for stat in data.speaker_stats:
                time_str = format_timestamp_readable(stat.speaking_time)
                lines.append(
                    f"{stat.speaker_name:<20} {stat.word_count:>8,} {time_str:>10} {stat.word_percent:>5.1f}%"
                )
            lines.append("")
            lines.append("-" * 50)
            lines.append("")

        # AI Summary
        if data.ai_summary:
            lines.append("AI Summary")
            lines.append("-" * 50)
            lines.append(data.ai_summary.get("summary", ""))
            lines.append("")

            key_points = data.ai_summary.get("key_points", [])
            if key_points:
                lines.append("Key Points:")
                for point in key_points:
                    lines.append(f"  • {point}")
                lines.append("")

            action_items = data.ai_summary.get("action_items", [])
            if action_items:
                lines.append("Action Items:")
                for item in action_items:
                    lines.append(f"  • {item}")
                lines.append("")

            topics = data.ai_summary.get("topics", [])
            if topics:
                lines.append(f"Topics: {', '.join(topics)}")
                lines.append("")

            lines.append("-" * 50)
            lines.append("")

        # Segments
        for segment in data.segments:
            speaker_name = data.speakers.get(segment.speaker or "", segment.speaker)
            speaker_prefix = f"[{speaker_name}] " if speaker_name else ""

            # Build annotation markers
            markers = []
            if segment.edited:
                markers.append("edited")
            if segment.highlight_color:
                markers.append(f"highlighted:{segment.highlight_color}")
            marker_str = f" [{', '.join(markers)}]" if markers else ""

            if include_timestamps:
                timestamp = format_timestamp_readable(segment.start_time)
                lines.append(f"[{timestamp}] {speaker_prefix}{segment.text}{marker_str}")
            else:
                lines.append(f"{speaker_prefix}{segment.text}{marker_str}")

            # Add comments if present
            if segment.comments:
                for comment in segment.comments:
                    lines.append(f"    > {comment}")

        return "\n".join(lines)

    def export_srt(self, data: ExportData) -> str:
        """Export transcript to SRT subtitle format.

        Args:
            data: Export data

        Returns:
            SRT content
        """
        lines = []

        for i, segment in enumerate(data.segments, 1):
            # Sequence number
            lines.append(str(i))

            # Timestamp
            start = format_timestamp_srt(segment.start_time)
            end = format_timestamp_srt(segment.end_time)
            lines.append(f"{start} --> {end}")

            # Text with optional speaker
            speaker_name = data.speakers.get(segment.speaker or "", segment.speaker)
            if speaker_name:
                lines.append(f"<v {speaker_name}>{segment.text}")
            else:
                lines.append(segment.text)

            # Blank line between entries
            lines.append("")

        return "\n".join(lines)

    def export_vtt(self, data: ExportData) -> str:
        """Export transcript to WebVTT subtitle format.

        Args:
            data: Export data

        Returns:
            VTT content
        """
        lines = ["WEBVTT", ""]

        # Metadata
        if data.title:
            lines.append(f"NOTE Title: {data.title}")
        if data.language:
            lines.append(f"NOTE Language: {data.language}")
        lines.append("")

        for i, segment in enumerate(data.segments, 1):
            # Add notes for highlights/edits/comments before the cue
            notes = []
            if segment.edited:
                notes.append("This segment was manually edited")
            if segment.highlight_color:
                notes.append(f"Highlighted: {segment.highlight_color}")
            if segment.comments:
                for comment in segment.comments:
                    notes.append(f"Comment: {comment}")
            if notes:
                lines.append(f"NOTE Segment {i} annotations:")
                for note in notes:
                    lines.append(f"  {note}")
                lines.append("")

            # Cue identifier (optional but useful)
            lines.append(f"cue-{i}")

            # Timestamp
            start = format_timestamp_vtt(segment.start_time)
            end = format_timestamp_vtt(segment.end_time)
            lines.append(f"{start} --> {end}")

            # Text with optional speaker
            speaker_name = data.speakers.get(segment.speaker or "", segment.speaker)
            if speaker_name:
                lines.append(f"<v {speaker_name}>{segment.text}")
            else:
                lines.append(segment.text)

            # Blank line between entries
            lines.append("")

        return "\n".join(lines)

    def export_docx(self, data: ExportData) -> bytes:
        """Export transcript to DOCX format.

        Args:
            data: Export data

        Returns:
            DOCX file bytes
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
        except ImportError:
            raise ImportError(
                "python-docx is not installed. Install with: pip install python-docx"
            )

        # Highlight color mapping
        HIGHLIGHT_COLORS = {
            "yellow": "FFFF00",
            "green": "90EE90",
            "blue": "87CEEB",
            "pink": "FFB6C1",
            "orange": "FFA500",
            "purple": "DDA0DD",
        }

        def set_highlight(run, color_name: str):
            """Set background highlight color on a run."""
            color_hex = HIGHLIGHT_COLORS.get(color_name.lower(), "FFFF00")
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), color_hex)
            run._r.get_or_add_rPr().append(shd)

        doc = Document()

        # Title
        title = doc.add_heading(data.title, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadata table
        doc.add_paragraph()
        table = doc.add_table(rows=0, cols=2)
        table.style = "Table Grid"

        if data.language:
            row = table.add_row()
            row.cells[0].text = "Language"
            row.cells[1].text = data.language.upper()

        if data.word_count:
            row = table.add_row()
            row.cells[0].text = "Word Count"
            row.cells[1].text = f"{data.word_count:,}"

        if data.duration_seconds:
            row = table.add_row()
            row.cells[0].text = "Duration"
            row.cells[1].text = format_timestamp_readable(data.duration_seconds)

        if data.model_used:
            row = table.add_row()
            row.cells[0].text = "Model"
            row.cells[1].text = data.model_used

        doc.add_paragraph()
        doc.add_heading("Transcript", level=1)

        # Segments
        current_speaker = None
        all_comments = []  # Collect comments for end section
        for segment in data.segments:
            speaker_name = data.speakers.get(segment.speaker or "", segment.speaker)

            # Add speaker heading when speaker changes
            if speaker_name and speaker_name != current_speaker:
                current_speaker = speaker_name
                speaker_para = doc.add_paragraph()
                speaker_run = speaker_para.add_run(f"{speaker_name}")
                speaker_run.bold = True
                speaker_run.font.size = Pt(11)

            # Add segment with timestamp
            timestamp = format_timestamp_readable(segment.start_time)
            para = doc.add_paragraph()

            # Timestamp in gray
            ts_run = para.add_run(f"[{timestamp}] ")
            ts_run.font.size = Pt(9)
            ts_run.font.color.rgb = RGBColor(128, 128, 128)

            # Text with highlight if present
            text_run = para.add_run(segment.text)
            text_run.font.size = Pt(11)
            if segment.highlight_color:
                set_highlight(text_run, segment.highlight_color)

            # Add edited marker
            if segment.edited:
                edited_run = para.add_run(" [edited]")
                edited_run.font.size = Pt(8)
                edited_run.font.color.rgb = RGBColor(128, 128, 128)
                edited_run.italic = True

            # Collect comments
            if segment.comments:
                for comment in segment.comments:
                    all_comments.append((timestamp, speaker_name or "Unknown", comment))

        # Add comments section if there are any
        if all_comments:
            doc.add_paragraph()
            doc.add_heading("Comments", level=1)
            for ts, speaker, comment in all_comments:
                para = doc.add_paragraph()
                ts_run = para.add_run(f"[{ts}] {speaker}: ")
                ts_run.bold = True
                ts_run.font.size = Pt(10)
                comment_run = para.add_run(comment)
                comment_run.font.size = Pt(10)

        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def export_pdf(self, data: ExportData) -> bytes:
        """Export transcript to PDF format.

        Args:
            data: Export data

        Returns:
            PDF file bytes
        """
        try:
            from reportlab.lib import colors
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
            from reportlab.lib.units import inch
            from reportlab.platypus import (
                Paragraph,
                SimpleDocTemplate,
                Spacer,
                Table,
                TableStyle,
            )
        except ImportError:
            raise ImportError(
                "reportlab is not installed. Install with: pip install reportlab"
            )

        # Highlight color mapping to hex
        HIGHLIGHT_COLORS = {
            "yellow": "#FFFF00",
            "green": "#90EE90",
            "blue": "#87CEEB",
            "pink": "#FFB6C1",
            "orange": "#FFA500",
            "purple": "#DDA0DD",
        }

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        elements = []

        # Title style
        title_style = ParagraphStyle(
            "Title",
            parent=styles["Heading1"],
            fontSize=18,
            alignment=1,  # Center
            spaceAfter=20,
        )

        # Speaker style
        speaker_style = ParagraphStyle(
            "Speaker",
            parent=styles["Normal"],
            fontSize=11,
            fontName="Helvetica-Bold",
            spaceBefore=12,
            spaceAfter=4,
        )

        # Segment style
        segment_style = ParagraphStyle(
            "Segment",
            parent=styles["Normal"],
            fontSize=10,
            leftIndent=20,
            spaceBefore=2,
            spaceAfter=2,
        )

        # Comment style
        comment_style = ParagraphStyle(
            "Comment",
            parent=styles["Normal"],
            fontSize=9,
            leftIndent=20,
            spaceBefore=4,
            spaceAfter=4,
        )

        # Timestamp style (unused but kept for reference)
        timestamp_style = ParagraphStyle(
            "Timestamp",
            parent=styles["Normal"],
            fontSize=8,
            textColor=colors.gray,
        )

        # Title
        elements.append(Paragraph(data.title, title_style))

        # Metadata table
        metadata = []
        if data.language:
            metadata.append(["Language", data.language.upper()])
        if data.word_count:
            metadata.append(["Word Count", f"{data.word_count:,}"])
        if data.duration_seconds:
            metadata.append(["Duration", format_timestamp_readable(data.duration_seconds)])
        if data.model_used:
            metadata.append(["Model", data.model_used])

        if metadata:
            t = Table(metadata, colWidths=[1.5 * inch, 3 * inch])
            t.setStyle(
                TableStyle([
                    ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, -1), 9),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ])
            )
            elements.append(t)
            elements.append(Spacer(1, 20))

        # Transcript heading
        elements.append(Paragraph("Transcript", styles["Heading2"]))
        elements.append(Spacer(1, 10))

        # Segments
        current_speaker = None
        all_comments = []  # Collect comments for end section
        for segment in data.segments:
            speaker_name = data.speakers.get(segment.speaker or "", segment.speaker)

            # Add speaker heading when speaker changes
            if speaker_name and speaker_name != current_speaker:
                current_speaker = speaker_name
                elements.append(Paragraph(speaker_name, speaker_style))

            # Add segment with timestamp
            timestamp = format_timestamp_readable(segment.start_time)

            # Build text with highlight background if present
            text_content = segment.text
            if segment.highlight_color:
                bg_color = HIGHLIGHT_COLORS.get(segment.highlight_color.lower(), "#FFFF00")
                text_content = f"<span backColor='{bg_color}'>{text_content}</span>"

            # Add edited marker
            edited_marker = ""
            if segment.edited:
                edited_marker = " <font color='gray' size='7'><i>[edited]</i></font>"

            text = f"<font color='gray' size='8'>[{timestamp}]</font> {text_content}{edited_marker}"
            elements.append(Paragraph(text, segment_style))

            # Collect comments
            if segment.comments:
                for comment in segment.comments:
                    all_comments.append((timestamp, speaker_name or "Unknown", comment))

        # Add comments section if there are any
        if all_comments:
            elements.append(Spacer(1, 20))
            elements.append(Paragraph("Comments", styles["Heading2"]))
            elements.append(Spacer(1, 10))
            for ts, speaker, comment in all_comments:
                text = f"<b>[{ts}] {speaker}:</b> {comment}"
                elements.append(Paragraph(text, comment_style))

        doc.build(elements)
        buffer.seek(0)
        return buffer.getvalue()


# Singleton instance
export_service = ExportService()
